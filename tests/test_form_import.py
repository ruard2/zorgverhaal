import io

from docx import Document

from app.form_import_service import extract_document_text, fidelity_errors, proposal_to_schema
from app.schemas import FormImportProposal


def proposal(label="Observaties"):
    return FormImportProposal.model_validate({
        "title": "Dagformulier",
        "purpose": "Dagelijkse registratie",
        "suggested_form_type": "dagformulier",
        "suggested_cadence": "daily",
        "sections": [{"title": "Rapportage", "fields": [{"id": "observations", "label": label, "type": "textarea", "required": True, "options": [], "source_quote": "Observaties *"}]}],
        "uncertainties": [],
        "fidelity_note": "Bron was leesbaar",
    })


def test_docx_paragraphs_and_tables_are_extracted_in_order():
    document = Document()
    document.add_heading("Dagformulier", level=1)
    document.add_paragraph("Rapportage")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Observaties *"
    table.cell(0, 1).text = "Vrije tekst"
    buffer = io.BytesIO(); document.save(buffer)
    text = extract_document_text(buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "formulier.docx")
    assert text.index("Dagformulier") < text.index("Rapportage") < text.index("Observaties * | Vrije tekst")


def test_fidelity_validation_rejects_changed_visible_label():
    source = "Dagformulier\nRapportage\nObservaties *"
    assert fidelity_errors(proposal(), source) == []
    errors = fidelity_errors(proposal("Waarnemingen"), source)
    assert any("Veldlabel" in error for error in errors)


def test_proposal_converts_without_ai_only_fields():
    schema = proposal_to_schema(proposal())
    field = schema["sections"][0]["fields"][0]
    assert field == {"id": "observations", "label": "Observaties", "type": "textarea", "required": True, "options": []}
