from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "Houvast-werkgeversdemo.docx"
PETROL = RGBColor(18, 92, 87)
INK = RGBColor(23, 59, 57)
MUTED = RGBColor(102, 123, 119)


def font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), size=14, bold=True, color=PETROL)


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    font(p.add_run(text))
    return p


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.78)
    section.left_margin = section.right_margin = Inches(0.9)
    section.header_distance = section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    font(kicker.add_run("WERKGEVERSDEMO"), size=9, bold=True, color=PETROL)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    font(title.add_run("Minder registratietijd. Sterkere zorgverslagen."), size=25, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(15)
    font(subtitle.add_run("Houvast helpt medewerkers sneller en vollediger rapporteren, terwijl de werkgever overzicht en controle houdt."), size=12, color=MUTED)

    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.16)
    callout.paragraph_format.right_indent = Inches(0.16)
    callout.paragraph_format.space_before = Pt(2)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.12
    shade(callout, "EDF5F0")
    font(callout.add_run("De winst: "), bold=True, color=PETROL)
    font(callout.add_run("minder losse vragen en dubbel werk, meer uniforme formulieren en sneller inzicht in wat aandacht nodig heeft."))

    add_heading(doc, "Zo werkt de demonstratie")
    steps = [
        "Log in met het verstrekte werkgeversaccount.",
        "Kies op het overzicht ‘Bekijk als demomedewerker’.",
        "Selecteer een cliënt en spreek of typ een zorgverslag.",
        "Bekijk wat automatisch is ingevuld. Alleen noodzakelijke vragen verschijnen, tegelijk.",
        "Controleer het concept en sla het definitief op.",
        "Kies ‘Terug naar beheer’. De rapportage staat direct in het cliëntdossier.",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.10
        font(p.add_run(step), size=10.5)

    add_heading(doc, "Wat de werkgever krijgt")
    benefits = [
        "Een actueel overzicht van rapportages, formulieren, zorgminuten en open acties.",
        "Persoonlijke medewerkersaccounts via een eenmalige uitnodigingslink.",
        "AI als hulpmiddel; de zorgmedewerker blijft verantwoordelijk voor controle en opslag.",
        "Definitieve rapportages blijven intact. Aanvullingen worden apart en herleidbaar vastgelegd.",
        "Organisatie-eigen cliënten, diensten en formulieren in één werkomgeving.",
    ]
    for item in benefits:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.10
        font(p.add_run(item), size=10.5)

    add_heading(doc, "Nieuwe medewerker uitnodigen")
    add_body(doc, "Open Beheer, kies ‘Nodig medewerker uit’ en vul naam en eventueel werk-e-mailadres in. Houvast maakt automatisch een medewerkernummer en persoonlijke link. Na activatie logt de medewerker voortaan in met het eigen e-mailadres en wachtwoord.", after=0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("Houvast | demonstratie voor werkgevers"), size=8.5, color=MUTED)

    doc.core_properties.title = "Houvast werkgeversdemo"
    doc.core_properties.subject = "Korte uitleg en demonstratieroute voor werkgevers"
    doc.core_properties.author = "Houvast"
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
