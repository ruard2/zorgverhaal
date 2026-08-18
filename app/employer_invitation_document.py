from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PETROL = RGBColor(18, 92, 87)
INK = RGBColor(23, 59, 57)
MUTED = RGBColor(102, 123, 119)
LOGO = Path(__file__).parent / "assets" / "communitytools-logo.png"


def set_font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_hyperlink(paragraph, text, url):
    relation_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "125C57")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    properties.extend([color, underline]); run.append(properties)
    text_node = OxmlElement("w:t"); text_node.text = text
    run.append(text_node); hyperlink.append(run); paragraph._p.append(hyperlink)


def heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(9)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    set_font(paragraph.add_run(text), size=13, bold=True, color=PETROL)


def body(doc, text, after=5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    set_font(paragraph.add_run(text))


def build_employer_invitation_document(organization_name: str, contact_name: str, email: str, invite_url: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.66)
    section.left_margin = section.right_margin = Inches(0.85)
    section.header_distance = section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.08

    logo = document.add_paragraph(); logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo.paragraph_format.space_after = Pt(5)
    picture = logo.add_run().add_picture(str(LOGO), width=Inches(5.7), height=Inches(1.35))
    source = picture._inline.graphic.graphicData.pic.blipFill
    crop = OxmlElement("a:srcRect")
    crop.set("l", "15000"); crop.set("r", "15000"); crop.set("t", "37500"); crop.set("b", "37500")
    source.insert(1, crop)

    kicker = document.add_paragraph(); kicker.paragraph_format.space_after = Pt(2)
    set_font(kicker.add_run("PERSOONLIJKE UITNODIGING | DEMO-ZORG"), size=8.5, bold=True, color=PETROL)
    title = document.add_paragraph(); title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run(f"Welkom bij Demo-Zorg, {contact_name}"), size=22, bold=True)
    subtitle = document.add_paragraph(); subtitle.paragraph_format.space_after = Pt(10)
    set_font(subtitle.add_run(f"Voor {organization_name}: ervaar hoe rapporteren eenvoudiger, sneller en beter controleerbaar wordt."), size=11.5, color=MUTED)

    heading(document, "1. Log in op uw werkgeversaccount")
    body(document, f"Uw account is al aangemaakt. Gebruik als inlogmailadres: {email}. Het tijdelijke wachtwoord is ‘verandermij’.", after=3)
    link_paragraph = document.add_paragraph(); link_paragraph.paragraph_format.space_after = Pt(7)
    add_hyperlink(link_paragraph, "Klik hier om in te loggen op uw werkgeversaccount", invite_url)
    body(document, "Na het inloggen vraagt Demo-Zorg u direct om een eigen wachtwoord van minimaal 12 tekens te kiezen. Daarna opent automatisch het werkgeversportaal. Later kunt u ook via de gewone homepage inloggen met hetzelfde e-mailadres en uw eigen wachtwoord. De persoonlijke link is zeven dagen geldig en kan één keer worden gebruikt.")

    heading(document, "2. Vul de demonstratie")
    body(document, "Na het inloggen opent het werkgeversportaal. Ga naar Beheer, open Demo-inhoud en kies ‘Demo-inhoud toevoegen’. Er worden fictieve cliënten, zorgdoelen en formulieren klaargezet. Gebruik geen echte cliëntgegevens.")

    heading(document, "3. Doorloop de belangrijkste route")
    steps = [
        "Kies ‘Bekijk als demomedewerker’ op het overzicht.",
        "Selecteer een cliënt en spreek of typ een zorgverslag.",
        "Controleer wat automatisch is ingevuld en sla het verslag op.",
        "Kies ‘Terug naar beheer’ en bekijk het resultaat in het cliëntdossier.",
        "Vraag eventueel een gerichte aanvulling; de medewerker krijgt die als taak.",
    ]
    for text in steps:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.48)
        paragraph.paragraph_format.first_line_indent = Inches(-0.23)
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(text), size=10)

    heading(document, "Wat levert het op?")
    body(document, "Minder registratietijd en losse vragen, uniforme formulieren, sneller werkgeversinzicht en een controleerbaar cliëntdossier. AI maakt een concept; de zorgmedewerker controleert en beslist altijd zelf wat wordt opgeslagen.", after=0)

    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Demo-Zorg | ontwikkeld door CommunityTools"), size=8, color=MUTED)
    document.core_properties.title = f"Uitnodiging Demo-Zorg - {organization_name}"
    document.core_properties.author = "CommunityTools"
    output = BytesIO(); document.save(output)
    return output.getvalue()
